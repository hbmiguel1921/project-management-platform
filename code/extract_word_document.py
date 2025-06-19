#!/usr/bin/env python3
"""
Extractor para documentos Word (.docx) 
Extraer el contenido completo del archivo de historias de usuario.
"""

import os
import sys
from pathlib import Path

def extract_docx_content(docx_path, output_path):
    """Extrae el contenido de un archivo .docx"""
    try:
        # Instalar la librería necesaria
        os.system("uv add -q python-docx")
        
        from docx import Document
        
        print(f"Leyendo documento: {docx_path}")
        
        # Abrir el documento
        doc = Document(docx_path)
        
        # Extraer todo el contenido
        content = []
        content.append("# Historias de Usuario - Plataforma de Gestión de Proyectos de Software\n")
        content.append(f"Extraído el: 2025-06-19 05:38:05\n")
        content.append("---\n\n")
        
        # Extraer párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Detectar títulos y formatear apropiadamente
                text = paragraph.text.strip()
                
                # Si parece un título principal
                if any(keyword in text.lower() for keyword in ['historia', 'requisito', 'funcionalidad', 'epic', 'feature']):
                    if len(text) < 100 and text.isupper() or text.startswith(('Historia', 'Epic', 'Feature', 'RF', 'RNF')):
                        content.append(f"## {text}\n\n")
                    else:
                        content.append(f"{text}\n\n")
                else:
                    content.append(f"{text}\n\n")
        
        # Extraer tablas si las hay
        if doc.tables:
            content.append("\n## Tablas del Documento\n\n")
            for i, table in enumerate(doc.tables):
                content.append(f"### Tabla {i+1}\n\n")
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    content.append("| " + " | ".join(row_data) + " |\n")
                content.append("\n")
        
        # Escribir el contenido extraído
        full_content = "".join(content)
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_content)
        
        print(f"Contenido extraído y guardado en: {output_path}")
        print(f"Tamaño del contenido: {len(full_content)} caracteres")
        
        # Mostrar un preview del contenido
        lines = full_content.split('\n')
        print("\n--- PREVIEW DEL CONTENIDO ---")
        for i, line in enumerate(lines[:50]):  # Primeras 50 líneas
            print(f"{i+1:3d}: {line}")
        if len(lines) > 50:
            print(f"... y {len(lines) - 50} líneas más")
        
        return True
        
    except Exception as e:
        print(f"Error al extraer el documento: {str(e)}")
        return False

if __name__ == "__main__":
    docx_path = "/workspace/user_input_files/Historia de usuario Software Gestion de Proyectos_2025-06-18.docx"
    output_path = "/workspace/docs/historias_usuario_completas.md"
    
    success = extract_docx_content(docx_path, output_path)
    if success:
        print("Extracción completada exitosamente")
    else:
        print("Error en la extracción")
        sys.exit(1)
